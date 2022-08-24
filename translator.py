import requests
import json
import docx

IAM_TOKEN = 'fskjdbckjsbdkjcbak12ej3irj20r'   # your token
folder_id = 'kflwlmfakekjio'                  # your id
target_language = 'ru'
en_doc = docx.Document("C:/Users/Ghost/Desktop/translation/MHOAG/1.docx")
all_paras = en_doc.paragraphs
ru_doc = docx.Document()


for para in all_paras:
    text = para.text
    body = {
        "targetLanguageCode": target_language,
        "texts": text,
        "folderId": folder_id,
    }

    headers = {
        "Content-Type": "application/json",
        "Authorization": "Bearer {0}".format(IAM_TOKEN)
    }

    response = requests.post('https://translate.api.cloud.yandex.net/translate/v2/translate',
                             headers=headers,
                             json=body
                             )

    json_load = json.loads(response.text)
    if list(json_load.keys())[0] != 'translations':
        pass
    else:
        ru_doc.add_paragraph(json_load['translations'][0]['text'])
ru_doc.save("C:/Users/Ghost/Desktop/translation/MHOAGRU/1.docx")
